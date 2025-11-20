# pdf_hybrid_converter_v18_kor.py
# 하이브리드 PDF->DOCX v18 (이미지 크기 원본 반영)
#
# [v18 개선사항]
# 1. 이미지 삽입 시 크기(Width)를 강제 고정하지 않고, PDF 원본 크기(BBox)를 계산하여 적용
# 2. "헤더/푸터 + 2단 본문" 구조 지원 유지
# 3. (Box Area) 텍스트 복구 및 글머리 기호 정리 유지
#
import sys, os, io, fitz  # pymupdf
import docx
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# MATLAB 연동용 UTF-8 설정
try:
    import sys as _sys, io as _io
    _sys.stdout = _io.TextIOWrapper(_sys.stdout.buffer, encoding='utf-8')
    _sys.stderr = _io.TextIOWrapper(_sys.stderr.buffer, encoding='utf-8')
except Exception:
    pass

def clamp(x, a, b): return max(a, min(b, x))

# --- [Helper] 글머리 기호 정리 및 스타일 반환 ---
def clean_text_and_style(text):
    bullets = ['•', '-', 'o', '·', '➢', '', 'v', '', '\uf0b7'] 
    clean_txt = text.strip()
    for b in bullets:
        if clean_txt.startswith(b):
            return clean_txt[len(b):].strip(), 'List Bullet'
    return clean_txt, None

# --- [Helper] 고급 2단 레이아웃 감지 (헤더/푸터 분리) ---
def detect_layout_and_split(page, text_blocks, page_width, page_height):
    center_x = page_width / 2
    margin = 15
    
    sorted_blocks = sorted(text_blocks, key=lambda b: b[1])
    if not sorted_blocks: return False, [], [], [], []

    crossing_indices = []
    for i, b in enumerate(sorted_blocks):
        x0, y0, x1, y1 = b[:4]
        if x0 < center_x - margin and x1 > center_x + margin:
            crossing_indices.append(i)
            
    if not crossing_indices:
        start_idx = 0; end_idx = len(sorted_blocks)
    else:
        gaps = []
        prev = -1
        for curr in crossing_indices + [len(sorted_blocks)]:
            if curr - prev - 1 > 0: gaps.append((prev + 1, curr))
            prev = curr
        if not gaps: return False, sorted_blocks, [], [], []
        
        gaps.sort(key=lambda x: x[1] - x[0], reverse=True)
        start_idx, end_idx = gaps[0]
        if (end_idx - start_idx) < 2: return False, sorted_blocks, [], [], []

    top_blocks = sorted_blocks[:start_idx]
    body_candidates = sorted_blocks[start_idx:end_idx]
    bottom_blocks = sorted_blocks[end_idx:]
    
    left_blocks = []
    right_blocks = []
    for b in body_candidates:
        x0, y0, x1, y1 = b[:4]
        if x1 < center_x + margin: left_blocks.append(b)
        elif x0 > center_x - margin: right_blocks.append(b)
        else: left_blocks.append(b)

    if len(left_blocks) == 0 or len(right_blocks) == 0:
        return False, sorted_blocks, [], [], []
        
    return True, top_blocks, left_blocks, right_blocks, bottom_blocks


# --- [엔진 1] 이미지 ---
def find_image_elements(page, doc):
    img_list = page.get_images(full=True)
    final = []
    for img in img_list:
        try:
            xref = img[0]
            if xref == 0: continue
            bbox = page.get_image_bbox(img)
            if not bbox: continue
            r = fitz.Rect(bbox)
            if not r.is_valid or r.is_empty: continue
            final.append({"type": "image", "bbox": r.irect, "data": doc.extract_image(xref)["image"]})
        except: pass
    return final

# --- [엔진 2] 표 ---
def find_tables_hybrid(page):
    tables = page.find_tables()
    final = []
    for tab in tables:
        if tab.bbox and fitz.Rect(tab.bbox).is_valid and tab.row_count > 1:
            ext = tab.extract()
            if ext: final.append({"type": "data_table", "bbox": fitz.Rect(tab.bbox).irect, "data": ext})
    return final

# --- [엔진 3] 박스 ---
def find_drawing_bboxes(page):
    drawings = page.get_drawings()
    rects = [d["rect"] for d in drawings if d["rect"].get_area() > 100]
    unique = []
    for r in rects:
        dup = False
        for u in unique:
            if u.intersects(r) and (u.intersect(r).get_area() / r.get_area() > 0.8):
                dup = True; break
        if not dup: unique.append(r)
    return [{"type": "box_table", "bbox": r.irect} for r in unique]

# --- [Helper] 이미지 삽입 함수 (크기 자동 계산) ---
def insert_image_auto_size(paragraph, img_obj):
    """ PDF의 bbox 너비를 인치로 변환하여 워드에 삽입 """
    try:
        x0, y0, x1, y1 = img_obj["bbox"]
        width_pt = x1 - x0
        width_in = width_pt / 72.0  # PDF point to inch conversion
        
        # 너무 크면(6인치 초과) 줄이고, 너무 작으면(0.2인치 미만) 최소값 유지
        width_in = max(0.2, min(width_in, 6.5))
        
        run = paragraph.add_run()
        run.add_picture(io.BytesIO(img_obj["data"]), width=Inches(width_in))
    except Exception as e:
        print(f"  [!] 이미지 삽입 오류: {e}")

# --- 메인 변환 ---
def convert_pdf_to_word_v18(pdf_path, word_path):
    print(f"--- PDF->Word v18 (이미지 크기 자동 최적화) ---")
    doc = fitz.open(pdf_path)
    docx_doc = docx.Document()
    
    for pnum, page in enumerate(doc):
        print(f"Page {pnum+1} 처리...")
        w, h = page.rect.width, page.rect.height
        
        images = find_image_elements(page, doc)
        tables = find_tables_hybrid(page)
        boxes = find_drawing_bboxes(page)
        
        filter_rects = [fitz.Rect(el["bbox"]) for el in images + tables + boxes]
        
        text_blocks = []
        for b in page.get_text("blocks"):
            r = fitz.Rect(b[:4])
            is_in = False
            for fr in filter_rects:
                if fr.contains(r): is_in = True; break
            if not is_in and b[4].strip():
                text_blocks.append(b)

        is_2col, top_b, left_b, right_b, bot_b = detect_layout_and_split(page, text_blocks, w, h)
        
        def render_blocks(block_list):
            block_list.sort(key=lambda b: b[1])
            for b in block_list:
                txt, style = clean_text_and_style(b[4])
                if txt:
                    p = docx_doc.add_paragraph(txt)
                    if style: p.style = style
        
        if is_2col:
            print("  -> 2단 레이아웃 적용")
            
            # Top (Header) + 이미지 처리
            # 상단 영역(헤더)에 있는 이미지만 골라내기
            header_limit_y = top_b[-1][3] if top_b else 100
            top_imgs = [img for img in images if img["bbox"][1] < header_limit_y]
            
            for img in top_imgs:
                p = docx_doc.add_paragraph()
                insert_image_auto_size(p, img) # [수정됨] 크기 자동 계산
            
            render_blocks(top_b)
            
            # Body (2 Columns)
            table = docx_doc.add_table(rows=1, cols=2)
            table.autofit = False
            
            cell_l = table.cell(0, 0)
            left_b.sort(key=lambda b: b[1])
            for b in left_b:
                txt, style = clean_text_and_style(b[4])
                if txt:
                    p = cell_l.add_paragraph(txt)
                    if style: p.style = style
            
            cell_r = table.cell(0, 1)
            right_b.sort(key=lambda b: b[1])
            for b in right_b:
                txt, style = clean_text_and_style(b[4])
                if txt:
                    p = cell_r.add_paragraph(txt)
                    if style: p.style = style
            
            docx_doc.add_paragraph()
            render_blocks(bot_b)

        else:
            # 1단 레이아웃
            all_items = []
            for img in images: all_items.append({'type':'img', 'obj':img, 'y':img['bbox'][1]})
            for tab in tables: all_items.append({'type':'tab', 'obj':tab, 'y':tab['bbox'][1]})
            for box in boxes:  all_items.append({'type':'box', 'obj':box, 'y':box['bbox'][1]})
            for txt in text_blocks: all_items.append({'type':'txt', 'obj':txt, 'y':txt[1]})
            all_items.sort(key=lambda x: x['y'])
            
            for item in all_items:
                if item['type'] == 'txt':
                    txt, style = clean_text_and_style(item['obj'][4])
                    if txt:
                        p = docx_doc.add_paragraph(txt)
                        if style: p.style = style
                elif item['type'] == 'img':
                    p = docx_doc.add_paragraph()
                    insert_image_auto_size(p, item['obj']) # [수정됨] 크기 자동 계산
                elif item['type'] == 'tab':
                    data = item['obj']['data']
                    if not data: continue
                    t = docx_doc.add_table(rows=len(data), cols=len(data[0]))
                    t.style = 'Table Grid'
                    for r, row in enumerate(data):
                        for c, val in enumerate(row):
                             if val: t.cell(r, c).text = str(val)
                    docx_doc.add_paragraph()
                elif item['type'] == 'box':
                    box_rect = fitz.Rect(item['obj']['bbox'])
                    box_text = page.get_text(clip=box_rect).strip()
                    if not box_text: continue
                    t = docx_doc.add_table(rows=1, cols=1)
                    t.style = 'Table Grid'
                    txt, style = clean_text_and_style(box_text)
                    cell = t.cell(0,0)
                    cell.text = ""
                    p = cell.paragraphs[0]; p.text = txt
                    if style: p.style = style
                    docx_doc.add_paragraph()

        if pnum < len(doc) - 1:
            docx_doc.add_page_break()

    docx_doc.save(word_path)
    print(f"완료: {word_path}")

if __name__ == "__main__":
    if len(sys.argv) < 3:
        print("Usage: python pdf_hybrid_converter_v18_kor.py <input.pdf> <output.docx>")
        sys.exit(1)
    convert_pdf_to_word_v18(sys.argv[1], sys.argv[2])