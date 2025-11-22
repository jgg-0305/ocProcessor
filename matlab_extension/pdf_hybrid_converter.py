# pdf_hybrid_converter_v18_kor.py
# 하이브리드 PDF->DOCX v18 (이미지 크기 원본 반영)
#
# [주요 기능 및 알고리즘 요약]
# 1. 하이브리드 객체 추출: 텍스트, 이미지, 표, 도형(Box)을 각각 별도 엔진으로 추출
# 2. 공간 분석(Spatial Analysis): 요소들의 좌표(BBox)를 분석하여 중복 제거 및 배치 순서 결정
# 3. 2단 레이아웃 감지: 문서 중앙을 기준으로 좌/우 텍스트 블록을 분리하여 Word 표(Table)로 구조화
# 4. 동적 이미지 리사이징: PDF 내부 좌표계(Point)를 Word 단위(Inch)로 변환하여 원본 비율 유지

import sys, os, io      # 시스템, 파일 경로, 바이트 입출력 처리를 위한 표준 라이브러리
import fitz             # PyMuPDF: PDF 파일을 읽고 객체를 추출하는 핵심 라이브러리
import docx             # python-docx: Word(.docx) 파일을 생성하고 편집하는 라이브러리
from docx.shared import Inches, Pt  # Word 문서에서 크기 단위(인치, 포인트)를 다루기 위함
from docx.enum.text import WD_ALIGN_PARAGRAPH # 문단 정렬(좌, 우, 중앙)을 위한 열거형 상수

# --- [환경 설정] 콘솔 출력 인코딩 호환성 처리 ---
# Windows 콘솔이나 MATLAB 등 외부 프로그램에서 파이썬 스크립트를 호출할 때,
# 한글 출력이 깨지는 현상(UnicodeEncodeError)을 방지하기 위한 안전장치입니다.
try:
    import sys as _sys, io as _io
    # 표준 출력(stdout)을 UTF-8 인코딩으로 강제 전환
    _sys.stdout = _io.TextIOWrapper(_sys.stdout.buffer, encoding='utf-8')
    # 표준 에러(stderr)도 UTF-8 인코딩으로 강제 전환
    _sys.stderr = _io.TextIOWrapper(_sys.stderr.buffer, encoding='utf-8')
except Exception:
    pass # 설정 실패 시 무시하고 진행 (일반적인 파이썬 환경에서는 문제없음)

# --- [유틸리티] 값 제한 함수 (Clamp) ---
# 입력값 x가 최소값 a보다 작으면 a를, 최대값 b보다 크면 b를 반환합니다.
# 이미지 크기가 너무 작거나 문서 폭을 넘지 않도록 제한할 때 사용됩니다.
def clamp(x, a, b): return max(a, min(b, x))

# --- [Helper] 텍스트 정리 및 스타일 감지 함수 ---
# PDF에서 추출한 텍스트 앞의 특수문자(글머리 기호)를 처리합니다.
def clean_text_and_style(text):
    # 제거할 글머리 기호 목록 (유니코드 포함)
    bullets = ['•', '-', 'o', '·', '➢', '', 'v', '', '\uf0b7'] 
    
    # 텍스트 앞뒤 공백 제거
    clean_txt = text.strip()
    
    # 글머리 기호 목록을 순회하며 검사
    for b in bullets:
        if clean_txt.startswith(b):
            # 기호가 발견되면 기호 길이만큼 잘라내고(slice), 
            # Word의 'List Bullet' 스타일을 적용하라는 신호를 반환
            return clean_txt[len(b):].strip(), 'List Bullet'
            
    # 기호가 없으면 원본 텍스트와 스타일 없음(None) 반환
    return clean_txt, None

# --- [핵심 알고리즘] 2단 레이아웃 감지 및 분할 함수 ---
# 페이지 내 텍스트 블록의 배치를 분석하여 (헤더 / 2단 본문 / 푸터) 구조인지 판단합니다.
def detect_layout_and_split(page, text_blocks, page_width, page_height):
    center_x = page_width / 2  # 페이지 가로 중앙 좌표 계산
    margin = 15                # 중앙선 침범 여부를 판단할 때 허용할 오차 범위
    
    # 1. 텍스트 블록들을 Y좌표(세로 위치) 기준으로 정렬 (문서의 읽는 순서)
    sorted_blocks = sorted(text_blocks, key=lambda b: b[1])
    
    # 블록이 하나도 없으면 분석 불가 -> False 반환
    if not sorted_blocks: return False, [], [], [], []

    # 2. "페이지를 가로지르는(Crossing)" 블록 찾기
    # 주로 헤더(제목)나 푸터(페이지 번호)가 여기에 해당합니다.
    crossing_indices = []
    for i, b in enumerate(sorted_blocks):
        x0, y0, x1, y1 = b[:4] # 블록의 좌상단(x0,y0), 우하단(x1,y1) 좌표
        
        # 블록의 왼쪽이 중앙보다 좌측이고, 오른쪽이 중앙보다 우측이면 가로지르는 것으로 판단
        if x0 < center_x - margin and x1 > center_x + margin:
            crossing_indices.append(i) # 해당 블록의 인덱스 저장
            
    # 3. 본문 영역(Body) 탐색
    # 가로지르는 블록들 사이의 빈 구간(Gap) 중 가장 큰 곳이 본문일 가능성이 높음
    if not crossing_indices:
        # 가로지르는 블록이 없으면 전체를 하나의 구간으로 설정
        start_idx = 0; end_idx = len(sorted_blocks)
    else:
        gaps = [] # (구간 시작 인덱스, 구간 끝 인덱스)를 저장할 리스트
        prev = -1
        # 가로지르는 블록들의 인덱스를 기준으로 사이사이 구간을 계산
        for curr in crossing_indices + [len(sorted_blocks)]:
            if curr - prev - 1 > 0: # 사이에 블록이 존재한다면
                gaps.append((prev + 1, curr))
            prev = curr
            
        # 구간이 하나도 없으면(모두 통짜 블록이면) 2단 분리 불가
        if not gaps: return False, sorted_blocks, [], [], []
        
        # 가장 긴 구간(텍스트 블록이 제일 많은 곳)을 본문으로 선택
        gaps.sort(key=lambda x: x[1] - x[0], reverse=True)
        start_idx, end_idx = gaps[0]
        
        # 본문 블록이 2개 미만이면 굳이 2단으로 나눌 필요 없음
        if (end_idx - start_idx) < 2: return False, sorted_blocks, [], [], []

    # 4. 영역별 블록 리스트 분할
    top_blocks = sorted_blocks[:start_idx]             # 상단 영역 (헤더)
    body_candidates = sorted_blocks[start_idx:end_idx] # 중간 영역 (본문 후보)
    bottom_blocks = sorted_blocks[end_idx:]            # 하단 영역 (푸터)
    
    # 5. 본문 영역을 좌/우 컬럼으로 분류 (Spatial Classification)
    left_blocks = []
    right_blocks = []
    for b in body_candidates:
        x0, y0, x1, y1 = b[:4]
        # 블록의 오른쪽 끝(x1)이 중앙선보다 왼쪽에 있으면 -> 왼쪽 컬럼
        if x1 < center_x + margin: left_blocks.append(b)
        # 블록의 왼쪽 시작(x0)이 중앙선보다 오른쪽에 있으면 -> 오른쪽 컬럼
        elif x0 > center_x - margin: right_blocks.append(b)
        # 애매하게 걸친 경우, 기본적으로 왼쪽으로 배정 (데이터 유실 방지)
        else: left_blocks.append(b)

    # 한쪽 컬럼이 완전히 비어있다면 2단 레이아웃이 아니라고 판단
    if len(left_blocks) == 0 or len(right_blocks) == 0:
        return False, sorted_blocks, [], [], []
        
    # 2단 레이아웃이 맞음 (True) 및 분할된 리스트 반환
    return True, top_blocks, left_blocks, right_blocks, bottom_blocks


# --- [엔진 1] 이미지 추출 함수 ---
def find_image_elements(page, doc):
    # 페이지 내 모든 이미지 리스트 가져오기
    img_list = page.get_images(full=True)
    final = []
    
    for img in img_list:
        try:
            xref = img[0] # 이미지의 내부 참조 ID
            if xref == 0: continue # 잘못된 참조면 패스
            
            # get_image_bbox: 이미지가 페이지 내에서 차지하는 좌표(Rect) 계산
            # 이 단계가 없으면 이미지가 어디에 위치하는지 알 수 없음
            bbox = page.get_image_bbox(img)
            if not bbox: continue
            
            # 좌표 객체를 PyMuPDF Rect로 변환
            r = fitz.Rect(bbox)
            # 유효하지 않거나 빈 영역이면 패스
            if not r.is_valid or r.is_empty: continue
            
            # 결과 리스트에 (타입, 좌표, 이미지 바이너리 데이터) 저장
            final.append({"type": "image", "bbox": r.irect, "data": doc.extract_image(xref)["image"]})
        except: pass # 오류 발생 시 해당 이미지는 건너뜀
    return final

# --- [엔진 2] 표(Table) 추출 함수 ---
def find_tables_hybrid(page):
    # PyMuPDF의 내장 표 감지 알고리즘 실행
    tables = page.find_tables()
    final = []
    
    for tab in tables:
        # 표의 좌표가 유효하고, 행(Row) 개수가 1개 초과인 경우만 추출 (단순 줄글 오인 방지)
        if tab.bbox and fitz.Rect(tab.bbox).is_valid and tab.row_count > 1:
            ext = tab.extract() # 표 안의 텍스트를 2차원 리스트로 추출
            if ext: 
                final.append({"type": "data_table", "bbox": fitz.Rect(tab.bbox).irect, "data": ext})
    return final

# --- [엔진 3] 박스(Drawing) 영역 추출 함수 ---
# 텍스트 상자나 강조를 위한 사각형 도형을 찾아냅니다.
def find_drawing_bboxes(page):
    # 페이지 내 벡터 그래픽(선, 도형 등) 가져오기
    drawings = page.get_drawings()
    # 면적이 100 이상인 사각형만 필터링 (너무 작은 점이나 선은 무시)
    rects = [d["rect"] for d in drawings if d["rect"].get_area() > 100]
    unique = []
    
    # 중복 영역 제거 (겹치는 박스가 여러 개 그려진 경우 하나로 합침)
    for r in rects:
        dup = False
        for u in unique:
            # 기존 박스(u)와 현재 박스(r)가 교차하고, 교차 면적이 80% 이상이면 중복으로 간주
            if u.intersects(r) and (u.intersect(r).get_area() / r.get_area() > 0.8):
                dup = True; break
        if not dup: unique.append(r)
        
    # 결과 포맷에 맞춰 리스트 반환
    return [{"type": "box_table", "bbox": r.irect} for r in unique]

# --- [Helper] 이미지 삽입 함수 (크기 자동 최적화) ---
# v18의 핵심 기능: PDF 상의 크기를 계산해 Word에 적절한 크기로 넣습니다.
def insert_image_auto_size(paragraph, img_obj):
    """ PDF의 bbox 너비를 인치로 변환하여 워드에 삽입 """
    try:
        # 이미지의 좌표 (좌, 상, 우, 하)
        x0, y0, x1, y1 = img_obj["bbox"]
        
        # PDF 단위(Point) 너비 계산
        width_pt = x1 - x0
        
        # Word 단위(Inch)로 변환 (1 inch = 72 points)
        width_in = width_pt / 72.0
        
        # 안전장치(Clamping):
        # 너무 작으면(0.2인치 미만) -> 0.2인치로 (보이게 하기 위해)
        # 너무 크면(6.5인치 초과) -> 6.5인치로 (문서 밖으로 튀어나감 방지)
        width_in = max(0.2, min(width_in, 6.5))
        
        # 현재 문단(paragraph)에 이미지 추가
        run = paragraph.add_run()
        # 바이트 스트림으로 이미지를 전달하고, 계산된 너비 지정
        run.add_picture(io.BytesIO(img_obj["data"]), width=Inches(width_in))
    except Exception as e:
        # 이미지 처리 중 에러가 나도 프로그램이 죽지 않고 로그만 출력
        print(f"  [!] 이미지 삽입 오류: {e}")

# --- [메인 로직] 전체 변환 프로세스 ---
def convert_pdf_to_word_v18(pdf_path, word_path):
    print(f"--- PDF->Word v18 (이미지 크기 자동 최적화) ---")
    
    # PDF 파일 열기
    doc = fitz.open(pdf_path)
    # 빈 Word 문서 생성
    docx_doc = docx.Document()
    
    # 각 페이지를 순회하며 처리
    for pnum, page in enumerate(doc):
        print(f"Page {pnum+1} 처리...")
        
        # 페이지의 전체 폭(w)과 높이(h) 가져오기
        w, h = page.rect.width, page.rect.height
        
        # 1. 3가지 엔진을 가동하여 시각 요소(이미지, 표, 박스) 먼저 추출
        images = find_image_elements(page, doc)
        tables = find_tables_hybrid(page)
        boxes = find_drawing_bboxes(page)
        
        # 2. 텍스트 추출 전, '중복 방지용' 영역 리스트 만들기
        # 표나 이미지 영역 안에 있는 텍스트는 일반 텍스트로 또 뽑히면 안 되므로 영역을 기억함
        filter_rects = [fitz.Rect(el["bbox"]) for el in images + tables + boxes]
        
        text_blocks = []
        # 페이지의 모든 텍스트 블록 가져오기
        for b in page.get_text("blocks"):
            r = fitz.Rect(b[:4]) # 현재 텍스트 블록의 좌표
            is_in = False
            
            # 현재 텍스트가 필터링 영역(이미지/표 등) 안에 포함되는지 검사
            for fr in filter_rects:
                if fr.contains(r): is_in = True; break
            
            # 포함되지 않고(순수 텍스트), 빈 문자열이 아니면 리스트에 추가
            if not is_in and b[4].strip():
                text_blocks.append(b)

        # 3. 레이아웃 분석 수행 (2단인지 1단인지 판단)
        is_2col, top_b, left_b, right_b, bot_b = detect_layout_and_split(page, text_blocks, w, h)
        
        # [내부 함수] 블록 리스트를 받아 Word 문단으로 찍어내는 역할
        def render_blocks(block_list):
            block_list.sort(key=lambda b: b[1]) # Y좌표 순으로 정렬
            for b in block_list:
                # 글머리 기호 정리 및 스타일 가져오기
                txt, style = clean_text_and_style(b[4])
                if txt:
                    p = docx_doc.add_paragraph(txt) # 문단 추가
                    if style: p.style = style       # 스타일 적용
        
        # [조건 분기 A] 2단 레이아웃인 경우
        if is_2col:
            print("  -> 2단 레이아웃 적용")
            
            # --- [Header 처리] ---
            # 헤더 영역의 경계선(Y좌표) 계산
            header_limit_y = top_b[-1][3] if top_b else 100
            
            # 전체 이미지 중 헤더 영역(위쪽)에 있는 것만 골라냄
            top_imgs = [img for img in images if img["bbox"][1] < header_limit_y]
            
            # 헤더 이미지 삽입
            for img in top_imgs:
                p = docx_doc.add_paragraph()
                insert_image_auto_size(p, img) # 크기 자동 계산
            
            # 헤더 텍스트 삽입
            render_blocks(top_b)
            
            # --- [Body 처리] ---
            # 2단 구성을 위해 Word 표(1행 2열) 생성
            table = docx_doc.add_table(rows=1, cols=2)
            table.autofit = False # 표 너비 자동 조절 끄기 (레이아웃 고정)
            
            # 왼쪽 칸(Cell 0,0)에 왼쪽 블록들 채우기
            cell_l = table.cell(0, 0)
            left_b.sort(key=lambda b: b[1])
            for b in left_b:
                txt, style = clean_text_and_style(b[4])
                if txt:
                    p = cell_l.add_paragraph(txt)
                    if style: p.style = style
            
            # 오른쪽 칸(Cell 0,1)에 오른쪽 블록들 채우기
            cell_r = table.cell(0, 1)
            right_b.sort(key=lambda b: b[1])
            for b in right_b:
                txt, style = clean_text_and_style(b[4])
                if txt:
                    p = cell_r.add_paragraph(txt)
                    if style: p.style = style
            
            # --- [Footer 처리] ---
            docx_doc.add_paragraph() # 표와 푸터 사이 간격
            render_blocks(bot_b)     # 푸터 텍스트 삽입

        # [조건 분기 B] 1단 레이아웃(일반)인 경우
        else:
            # --- [통합 배치] ---
            # 모든 요소(텍스트, 이미지, 표, 박스)를 하나의 리스트에 담음
            # 이때 각 요소의 Y좌표('y')를 함께 저장하여 위치를 기록
            all_items = []
            for img in images: all_items.append({'type':'img', 'obj':img, 'y':img['bbox'][1]})
            for tab in tables: all_items.append({'type':'tab', 'obj':tab, 'y':tab['bbox'][1]})
            for box in boxes:  all_items.append({'type':'box', 'obj':box, 'y':box['bbox'][1]})
            for txt in text_blocks: all_items.append({'type':'txt', 'obj':txt, 'y':txt[1]})
            
            # Y좌표(세로 위치) 기준으로 오름차순 정렬 -> 문서의 위에서 아래 순서대로 배치됨
            all_items.sort(key=lambda x: x['y'])
            
            # 정렬된 순서대로 하나씩 꺼내서 Word에 기록
            for item in all_items:
                # 1. 텍스트인 경우
                if item['type'] == 'txt':
                    txt, style = clean_text_and_style(item['obj'][4])
                    if txt:
                        p = docx_doc.add_paragraph(txt)
                        if style: p.style = style
                
                # 2. 이미지인 경우
                elif item['type'] == 'img':
                    p = docx_doc.add_paragraph()
                    insert_image_auto_size(p, item['obj']) # 크기 자동 계산 적용
                
                # 3. 표(Table)인 경우
                elif item['type'] == 'tab':
                    data = item['obj']['data'] # 표 데이터 가져오기
                    if not data: continue
                    # 데이터 크기에 맞는 Word 표 생성
                    t = docx_doc.add_table(rows=len(data), cols=len(data[0]))
                    t.style = 'Table Grid' # 기본 격자 테두리 스타일
                    
                    # 셀 하나하나 채워 넣기
                    for r, row in enumerate(data):
                        for c, val in enumerate(row):
                             if val: t.cell(r, c).text = str(val)
                    docx_doc.add_paragraph() # 표 아래 빈 줄 추가
                
                # 4. 박스(도형)인 경우
                elif item['type'] == 'box':
                    # 박스 영역 좌표를 이용해 해당 영역 안의 텍스트를 다시 긁어옴(Clipping)
                    box_rect = fitz.Rect(item['obj']['bbox'])
                    box_text = page.get_text(clip=box_rect).strip()
                    if not box_text: continue # 빈 박스면 패스
                    
                    # 1행 1열짜리 표를 만들어 박스(테두리) 효과를 냄
                    t = docx_doc.add_table(rows=1, cols=1)
                    t.style = 'Table Grid'
                    
                    # 내용 채우기
                    txt, style = clean_text_and_style(box_text)
                    cell = t.cell(0,0)
                    cell.text = "" # 초기화
                    p = cell.paragraphs[0]; p.text = txt
                    if style: p.style = style
                    docx_doc.add_paragraph()

        # 마지막 페이지가 아니라면 '페이지 나누기(Page Break)' 추가
        if pnum < len(doc) - 1:
            docx_doc.add_page_break()

    # 모든 처리가 끝나면 Word 파일로 저장
    docx_doc.save(word_path)
    print(f"완료: {word_path}")

# --- [실행 진입점] ---
if __name__ == "__main__":
    # 커맨드라인 인자 개수 확인 (스크립트명, 입력파일, 출력파일 총 3개 필요)
    if len(sys.argv) < 3:
        print("Usage: python pdf_hybrid_converter_v18_kor.py <input.pdf> <output.docx>")
        sys.exit(1) # 인자 부족 시 프로그램 종료
        
    # 변환 함수 실행
    convert_pdf_to_word_v18(sys.argv[1], sys.argv[2])


